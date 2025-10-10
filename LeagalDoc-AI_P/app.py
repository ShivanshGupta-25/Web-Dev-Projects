import os
import google.generativeai as genai
from flask import Flask, render_template, request, jsonify, redirect, url_for, session
from werkzeug.utils import secure_filename
import PyPDF2
import docx
import markdown  # pip install markdown
from flask_sqlalchemy import SQLAlchemy
from datetime import datetime

# ---- Flask setup ----
app = Flask(__name__)
app.config["UPLOAD_FOLDER"] = "uploads"
os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)
# ---- Database setup ----
app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///legal_assistant.db"  # SQLite file
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False

db = SQLAlchemy(app)

# Use a secret key for session usage
app.secret_key = os.getenv("FLASK_SECRET_KEY", "dev-secret-key")

# ---- Google Generative AI setup ----
# Make sure you have set the environment variable GOOGLE_API_KEY
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
model = genai.GenerativeModel("gemini-1.5-flash")

# Simple in-memory store for demo (not suitable for production)
DOCUMENT_CONTEXT = ""


# ---- Helpers ----
def extract_text(filepath: str) -> str:
    """Extract text from PDF, DOCX or TXT. Returns an error string for unsupported types."""
    lower = filepath.lower()
    if lower.endswith(".pdf"):
        text = ""
        with open(filepath, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text.strip()

    if lower.endswith(".docx"):
        doc = docx.Document(filepath)
        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

    if lower.endswith(".txt"):
        with open(filepath, "r", encoding="utf-8") as f:
            return f.read()

    return "Unsupported file format."


def target_language_name(lang_code: str) -> str:
    """Return a human-friendly language name from the code."""
    if lang_code == "en":
        return "English"
    if lang_code == "hi":
        return "Hindi"
    return lang_code  # fallback


def extract_response_text(resp) -> str:
    """
    Safely extract text from a model response object.
    SDKs differ in return shape; check common attributes and fallback to str(resp).
    """
    if resp is None:
        return ""
    # direct text attribute
    if hasattr(resp, "text"):
        return resp.text
    # some SDKs return candidates
    try:
        if hasattr(resp, "candidates") and resp.candidates:
            cand = resp.candidates[0]
            if hasattr(cand, "content"):
                return cand.content
            if hasattr(cand, "text"):
                return cand.text
            if isinstance(cand, dict):
                for k in ("content", "text", "output"):
                    if k in cand:
                        return cand[k]
    except Exception:
        pass
    # final fallback
    try:
        return str(resp)
    except Exception:
        return ""


# ---- Routes ----
@app.route("/")
def index():
    return render_template("index.html")


@app.route("/upload", methods=["POST"])
def upload():
    """
    Upload endpoint:
    - saves file
    - extracts text
    - stores a trimmed DOCUMENT_CONTEXT for later Q&A
    - generates a structured summary in the chosen language (en/hi)
    """
    global DOCUMENT_CONTEXT

    if "file" not in request.files:
        return redirect(url_for("index"))

    file = request.files["file"]
    if file.filename == "":
        return redirect(url_for("index"))

    # chosen language for summary (applies only to summarization)
    language = request.form.get("language", "en")
    session["language"] = language
    target_lang = target_language_name(language)

    # save uploaded file
    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config["UPLOAD_FOLDER"], filename)
    file.save(filepath)

    raw_text = extract_text(filepath)
    if not isinstance(raw_text, str) or raw_text.startswith("Unsupported"):
        return render_template(
            "result.html",
            original="",
            simplified=f"⚠️ {raw_text}",
            chosen_language=target_lang
        )

    # keep a limited document context for demo (avoid sending huge texts)
    DOCUMENT_CONTEXT = raw_text[:8000]

    # structured summarization prompt
    prompt = (
        f"You are a legal assistant. Read the legal text below and produce a clear, plain-language summary in {target_lang} only.\n\n"
        f"Formatting rules:\n"
        f"- Use clear section headings (Obligations, Rights, Risks, Important Dates, etc.).\n"
        f"- Use bullet points for details.\n"
        f"- Keep language simple and avoid legal jargon.\n"
        f"- Do NOT add extra commentary outside the summary.\n\n"
        f"Text:\n{DOCUMENT_CONTEXT}"
    )

    simplified_text_raw = ""
    try:
        resp = model.generate_content(prompt)
        simplified_text_raw = extract_response_text(resp) or "⚠️ AI did not return any text."
        # convert markdown-like text to HTML for nice rendering
        simplified_text_html = markdown.markdown(simplified_text_raw)
    except Exception as e:
        simplified_text_html = f"⚠️ Error while generating summary: {str(e)}"

    # Save document to DB (for demo, single user id = 1)
    user = User.query.first()
    if not user:
        user = User(username="demo_user")
        db.session.add(user)
        db.session.commit()

    new_doc = Document(
        filename=filename,
        language=language,
        original_text=DOCUMENT_CONTEXT,
        simplified_text=simplified_text_raw,
        user_id=user.id,
    )
    db.session.add(new_doc)
    db.session.commit()

    # Store doc_id in session to use for Q&A
    session["last_doc_id"] = new_doc.id

    return render_template(
        "result.html",
        original=raw_text[:2000],
        simplified=simplified_text_html,
        chosen_language=target_lang
    )


@app.route("/ask", methods=["POST"])
def ask():
    """
    Chat endpoint:
    - Accepts JSON { "question": "...", "chat_language": "auto"|"en"|"hi" }.
    - If chat_language == "auto", instruct the model to detect the question language and answer in it.
    - If chat_language is "en" or "hi", force the model to answer in that language.
    """
    global DOCUMENT_CONTEXT

    data = request.get_json(silent=True) or {}
    user_question = (data.get("question") or "").strip()
    # accept either key name for safety
    chat_lang = (data.get("chat_language") or data.get("chatLang") or "auto")

    # Check if document exists
    doc_id = session.get("last_doc_id")
    if not doc_id:
        return jsonify({"answer": "⚠️ Please upload a legal document first."})

    doc = Document.query.get(doc_id)
    if not doc:
        return jsonify({"answer": "⚠️ Document not found."})

    # if not DOCUMENT_CONTEXT:
    #     return jsonify({"answer": "⚠️ Please upload a legal document first."})

    # Build prompt depending on chat language selection
    if chat_lang == "auto":
        qa_prompt = (
            "You are a helpful AI legal assistant.\n"
            "Use the document context below to answer the user's question.\n"
            "RULE: Detect the language of the question and always respond in the SAME language as the question.\n"
            "Be concise and easy to understand.\n\n"
            f"Document:\n{DOCUMENT_CONTEXT}\n\n"
            f"Question: {user_question}\n"
            "Answer:"
        )
    else:
        # force target language
        target = target_language_name(chat_lang)
        qa_prompt = (
            f"You are a helpful AI legal assistant.\n"
            f"Use the document context below to answer the user's question in {target} only.\n"
            "Be concise and easy to understand.\n\n"
            f"Document:\n{DOCUMENT_CONTEXT}\n\n"
            f"Question: {user_question}\n"
            "Answer:"
        )

    try:
        resp = model.generate_content(qa_prompt)
        answer = extract_response_text(resp) or "⚠️ AI did not return any text."
    except Exception as e:
        answer = f"⚠️ Error while generating answer: {str(e)}"

    return jsonify({"answer": answer})


# Models ------------------------
class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class Document(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(200), nullable=False)
    language = db.Column(db.String(10), nullable=False, default="en")
    original_text = db.Column(db.Text, nullable=False)
    simplified_text = db.Column(db.Text, nullable=False)
    uploaded_at = db.Column(db.DateTime, default=datetime.utcnow)
    user_id = db.Column(db.Integer, db.ForeignKey("user.id"))

    user = db.relationship("User", backref=db.backref("documents", lazy=True))

with app.app_context():
    db.create_all()

# Run the app
if __name__ == "__main__":
    # host 0.0.0.0 is typical for containerized deployments; change if you want.
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=True)
